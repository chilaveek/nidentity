from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[2]
OUTPUT_DIR = ROOT / "output" / "doc"
DOCX_PATH = OUTPUT_DIR / "otchet_lab5_nidshop.docx"


def set_font(run, size=14, bold=False, italic=False):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    set_font(run)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(end)


def set_shading(cell, fill="F2F2F2"):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def configure_document(doc):
    section = doc.sections[0]
    section.left_margin = Cm(3)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    style.font.size = Pt(14)
    fmt = style.paragraph_format
    fmt.line_spacing = 1.5
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    fmt.first_line_indent = Cm(1.25)

    footer = section.footer
    footer.is_linked_to_previous = False
    footer.paragraphs[0].clear()
    add_page_number(footer.paragraphs[0])
    section.different_first_page_header_footer = True


def paragraph(doc, text, align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_indent=True, bold=False, italic=False, space_after=0):
    p = doc.add_paragraph()
    p.alignment = align
    fmt = p.paragraph_format
    fmt.line_spacing = 1.5
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(space_after)
    fmt.first_line_indent = Cm(1.25 if first_indent else 0)
    run = p.add_run(text)
    set_font(run, bold=bold, italic=italic)
    return p


def title_page(doc):
    section = doc.sections[0]
    section.different_first_page_header_footer = True

    for _ in range(4):
        paragraph(doc, "", align=WD_ALIGN_PARAGRAPH.CENTER, first_indent=False)

    paragraph(
        doc,
        "ОТЧЁТ ПО ЛАБОРАТОРНОЙ РАБОТЕ № 5",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        first_indent=False,
        bold=True,
    )
    paragraph(
        doc,
        "по дисциплине «Конструирование интернет-приложений»",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        first_indent=False,
    )
    paragraph(doc, "", align=WD_ALIGN_PARAGRAPH.CENTER, first_indent=False)
    paragraph(
        doc,
        "Тема: создание и просмотр заявок в веб-приложении",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        first_indent=False,
    )

    for _ in range(7):
        paragraph(doc, "", align=WD_ALIGN_PARAGRAPH.CENTER, first_indent=False)

    paragraph(doc, "Выполнил студент: ________________________________", first_indent=False)
    paragraph(doc, "Группа: ________________________________", first_indent=False)
    paragraph(doc, "Проверил: ________________________________", first_indent=False)

    for _ in range(6):
        paragraph(doc, "", align=WD_ALIGN_PARAGRAPH.CENTER, first_indent=False)

    paragraph(doc, "2026", align=WD_ALIGN_PARAGRAPH.CENTER, first_indent=False)


def heading(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    fmt = p.paragraph_format
    fmt.first_line_indent = Cm(0)
    fmt.line_spacing = 1.5
    fmt.space_before = Pt(6)
    fmt.space_after = Pt(6)
    run = p.add_run(text)
    set_font(run, bold=True)
    return p


def figure_placeholder(doc, num, caption, note, height_cm=6.3):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    row = table.rows[0]
    row.height = Cm(height_cm)
    row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    cell = row.cells[0]
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_shading(cell)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(note)
    set_font(run, size=12, italic=True)

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.first_line_indent = Cm(0)
    cap.paragraph_format.line_spacing = 1.5
    run = cap.add_run(f"Рис. {num}. {caption}")
    set_font(run, size=12)


def make_table(doc, title, rows, widths):
    paragraph(doc, title, align=WD_ALIGN_PARAGRAPH.LEFT, first_indent=False)
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx]
        for c_idx, value in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.width = Cm(widths[c_idx])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if r_idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.line_spacing = 1.0
            run = p.add_run(value)
            set_font(run, size=12, bold=(r_idx == 0))
        if r_idx == 0:
            for cell in row.cells:
                set_shading(cell, "D9E2F3")


def main():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    doc = Document()
    configure_document(doc)
    title_page(doc)

    doc.add_page_break()

    heading(doc, "Цель работы:")
    paragraph(
        doc,
        "Реализовать основной домен заявок в веб-приложении: создание обращения, "
        "просмотр списка обращений и открытие переписки по конкретной заявке с сохранением данных на сервере и отображением на клиенте.",
        first_indent=False,
    )

    heading(doc, "Ход работы")
    paragraph(
        doc,
        "В качестве функционального аналога модуля заявок использован раздел обращений в поддержку проекта "
        "NIDSHOP. На сервере присутствует сущность SupportTicket с полями Id, UserId, Subject, Status и CreatedAt. "
        "Для хранения переписки используется связанная сущность SupportMessage с текстом сообщения, отправителем, "
        "признаком сообщения администратора и временем создания. Такая модель покрывает ключевые требования лабораторной работы: "
        "есть сама заявка, ее автор, текущее состояние и подробные данные для просмотра.",
    )
    figure_placeholder(
        doc,
        "1",
        "Сущности SupportTicket и SupportMessage",
        "Место для скриншота моделей заявок и сообщений из серверной части проекта",
    )

    paragraph(
        doc,
        "Хранение и выборка заявок настроены через ApplicationDbContext. В контекст добавлены наборы данных "
        "SupportTickets и SupportMessages, а в конфигурации сущностей описаны связи с пользователем, индексы "
        "и правила удаления. За счет этого каждая заявка связывается с конкретным пользователем, а история "
        "сообщений может быть загружена отдельно при открытии деталей.",
    )

    paragraph(
        doc,
        "Работа с API сосредоточена в AuthController. Для пользовательского сценария реализованы четыре основных "
        "маршрута: POST /api/auth/support/tickets для создания заявки, GET /api/auth/support/tickets для получения "
        "списка собственных обращений, GET /api/auth/support/tickets/{ticketId}/messages для просмотра переписки "
        "по выбранной заявке и POST /api/auth/support/tickets/{ticketId}/message для отправки нового сообщения. "
        "Все методы требуют авторизацию. При обращении к данным сервер дополнительно проверяет, что пользователь "
        "запрашивает только свои заявки.",
    )
    make_table(
        doc,
        "Таблица 1 - Основные методы API для работы с заявками",
        [
            ["Метод", "Маршрут", "Назначение"],
            ["POST", "/api/auth/support/tickets", "Создание новой заявки"],
            ["GET", "/api/auth/support/tickets", "Получение списка заявок текущего пользователя"],
            ["GET", "/api/auth/support/tickets/{ticketId}/messages", "Просмотр сообщений по заявке"],
            ["POST", "/api/auth/support/tickets/{ticketId}/message", "Отправка сообщения в открытую заявку"],
        ],
        [2.1, 7.2, 7.2],
    )

    paragraph(
        doc,
        "При создании заявки сервер формирует новую запись SupportTicket и сразу добавляет первое сообщение в "
        "SupportMessage. Это соответствует реальному сценарию работы с обращениями: пользователь указывает тему, "
        "описывает проблему, после чего видит заявку в списке и может продолжить переписку без повторного создания записи.",
    )
    figure_placeholder(
        doc,
        "2",
        "Ответ сервера на создание новой заявки",
        "Место для скриншота вкладки Network или Swagger с успешным ответом POST-запроса",
    )

    paragraph(
        doc,
        "На клиентской стороне модуль заявок реализован на странице SupportPage. В режиме списка пользователь "
        "видит все свои обращения, их темы, дату создания, статус и количество сообщений. Для создания новой заявки "
        "открывается модальное окно с двумя полями: тема обращения и текст первого сообщения. После успешной отправки "
        "список обновляется и новое обращение сразу становится доступно для просмотра.",
    )
    figure_placeholder(
        doc,
        "3",
        "Страница списка обращений пользователя",
        "Место для скриншота страницы со списком заявок и кнопкой создания нового обращения",
    )
    figure_placeholder(
        doc,
        "4",
        "Модальное окно создания заявки",
        "Место для скриншота формы создания обращения с полями темы и текста сообщения",
    )

    paragraph(
        doc,
        "Если пользователь открывает конкретную заявку, интерфейс переключается в режим деталей. "
        "На странице отображаются тема, статус, время создания и вся история сообщений в формате чата. "
        "Сообщения администратора и пользователя визуально разделены, а в нижней части страницы доступна форма "
        "для отправки нового сообщения. Такой интерфейс выполняет роль страницы деталей заявки и дополняет базовый "
        "сценарий просмотра более удобной перепиской.",
    )
    figure_placeholder(
        doc,
        "5",
        "Страница деталей заявки и переписки",
        "Место для скриншота открытой заявки с историей сообщений и формой ответа",
    )

    paragraph(
        doc,
        "Дополнительно в клиентской реализации поддержан сценарий повторного открытия закрытой заявки. "
        "Если пользователь отправляет сообщение в ранее закрытое обращение, сервер автоматически меняет статус "
        "на Open. Это делает модуль более удобным и показывает, что логика состояния может обрабатываться не только "
        "в интерфейсе, но и на стороне сервера.",
    )
    figure_placeholder(
        doc,
        "6",
        "Запрос на получение списка заявок в инструментах разработчика браузера",
        "Место для скриншота вкладки Network с выбранным GET-запросом и JSON-ответом сервера",
    )

    heading(doc, "Вывод:")
    paragraph(
        doc,
        "В ходе выполнения лабораторной работы был рассмотрен и оформлен модуль заявок на основе раздела "
        "обращений в поддержку проекта NIDSHOP. На сервере реализовано хранение заявок и сообщений, настроены "
        "маршруты для создания и просмотра обращений, а на клиенте добавлены список заявок, форма создания и "
        "страница деталей с перепиской. В результате получен рабочий пользовательский сценарий, в котором заявка "
        "создается, сохраняется в базе данных, загружается через API и отображается в интерфейсе.",
    )

    doc.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    main()
